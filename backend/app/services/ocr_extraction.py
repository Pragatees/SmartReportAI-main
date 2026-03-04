import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import io
import os
from docx import Document

# Set Tesseract path (change if different in your system)
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


def extract_text_from_file(contents: bytes, filename: str):
    file_extension = os.path.splitext(filename)[1].lower()
    text = ""

    if file_extension == ".pdf":
        doc = fitz.open(stream=contents, filetype="pdf")

        for page_num in range(doc.page_count):
            page = doc[page_num]
            page_text = page.get_text()

            if page_text.strip():
                text += page_text + "\n"
            else:
                img_list = page.get_images(full=True)

                for img in img_list:
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]

                    img_obj = Image.open(io.BytesIO(image_bytes))
                    img_text = pytesseract.image_to_string(img_obj)

                    text += img_text + "\n"

            text += f"\n--- End of Page {page_num + 1} ---\n"

        doc.close()

    elif file_extension in [".png", ".jpg", ".jpeg", ".tif", ".tiff"]:
        img = Image.open(io.BytesIO(contents))
        text = pytesseract.image_to_string(img)

    elif file_extension in [".doc", ".docx"]:
        doc = Document(io.BytesIO(contents))

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"

    else:
        raise ValueError(f"Unsupported file format: {file_extension}")

    return text